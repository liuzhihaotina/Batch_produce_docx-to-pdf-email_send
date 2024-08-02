# pip install openpyxl
from numpy import iinfo
import pandas
#from docx import Document
import docx
# docx = docx.Document()
def replace_text(doc, old_text, new_text):
    # 遍历每个段落
    for p in doc.paragraphs:
        # 如果要搜索的内容在该段落
        if old_text in p.text:
            # 使用 runs 替换内容但不改变样式
            # 注意！runs 会根据样式分隔内容，确保被替换内容的样式一致
            for run in p.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)

if __name__ == "__main__":

    # majors = ['机械','电气']
    # majors = ['控制','动力']
    # 手动更改专业名称
    majors = ['动力']
    # 相对路径，同目录下
    oriData = '动力夏令营成绩邮箱.xlsx'

    for major in majors:
        df = pandas.read_excel(f'{oriData}')
        # print(df)

        # # 显示每段的内容
        # for p in document.paragraphs:
        #     print(p.text)

        # print(len(df))
        # exit()
        for i in range(len(df)):
            
            print(df.loc[i]['姓名'])
            name = df.loc[i]['姓名']
            rank = df.loc[i]['序号']
            # applied_major = df.loc[i]['拟申请学科名称']
            applied_major='动力'
            mark = str(df.loc[i]['面试分数'])
            print(mark)

            document = docx.Document(f'muban.docx')
            replace_text(document, '姓名', name)
            replace_text(document, '分数', mark)    # BUG：这里加了<>就不行，后面又可以了，可能是因为程序没保存
            replace_text(document, '专业', applied_major)

            # if df.loc[i]['类型'] == "学术型":
            #     replace_text(document, '<类型>', '学术型')
            # elif df.loc[i]['类型'] == "专业型":
            #     replace_text(document, '<类型>', '专业学位')
            # else:
            #     print('错误！未知类型。')
            #     exit()

            # directory = 'results/docx/'+major+'/'

            # 绝对路径
            directory ='D:/控制面试/Docx-20230727/Docx/results/docx/'+major+'/'
            filename = major+'-'+str(rank)+'-'+name+'.docx'
            document.save(directory+filename)
                # 哈尔滨工业大学（深圳）推免面试成绩单-姓名
        