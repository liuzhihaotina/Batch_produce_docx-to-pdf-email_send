# pip install openpyxl
from numpy import iinfo
import pandas
#from docx import Document
import docx  # 若安装的是docx先卸载掉，下载python-docx
import time
# docx = docx.Document()
def replace_text(doc, old_text, new_text):
    # 遍历每个段落
    for p in doc.paragraphs:
        # 如果要搜索的内容在该段落
        if old_text in p.text:
            # 使用 runs 替换内容但不改变样式
            # 注意！runs 会根据样式分隔内容，确保被替换内容的样式一致
            for run in p.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)

if __name__ == "__main__":

    # majors = ['机械','电气']
    # majors = ['控制','动力']
    # 手动更改专业名称
    majors = ['控制']
    # 相对路径
    oriData = f'./邮箱表格/{majors[0]}学科预录取函.xlsx'
    # 日期
    t = time.localtime()
    date=f'{t.tm_year}年{t.tm_mon}月{t.tm_mday}日'
    for major in majors:
        df = pandas.read_excel(f'{oriData}')

        for i in range(len(df)):
            print(df.loc[i]['姓名'])
            name = df.loc[i]['姓名']#综合成绩、证件号码
            # grade=df.loc[i]['综合成绩']
            num=str(df.loc[i]['证件号码'])
            rank = df.loc[i]['序号']
            lei=df.loc[i]['学术型/专业学位']
            m=df.loc[i]['专业和研究方向']
            # applied_major = df.loc[i]['拟申请学科名称']
            # applied_major='控制'
            mark = str(df.loc[i]['综合成绩'])
            # print(mark)
            if lei=='学术型':
                document = docx.Document(f'./word模板/学硕.docx')
            elif lei=='专业学位':
                document = docx.Document(f'./word模板/专硕.docx')
            else: # 直博生另外发送，这里不管
                continue
            replace_text(document, '姓名', name)
            replace_text(document, '分数', mark)
            replace_text(document, 'XX', num)
            replace_text(document, '方向', m)
            # 相对路径
            directory ='./生成文档/word/'+major+'/'
            filename = '2025年哈工大（深圳）机电学院推免生预接收函'+'-'+name+'.docx'
            document.save(directory+filename)
        